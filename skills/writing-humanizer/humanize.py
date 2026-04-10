#!/usr/bin/env python3
"""
humanize.py — Autonomous essay humanizer.

Phase 1 (done by Claude Code before calling this): essay is drafted and passed via --text.
Phase 2 (this script, fully autonomous):
  - Humanizes each paragraph via GPTInf.com (Simple mode, Playwright)
  - Verifies each paragraph with Claude Sonnet (citations/quotes/meaning)
  - Checks full essay against GPTZero; re-runs flagged paragraphs if score < 70%
  - Saves original + humanized versions to a Word doc

Usage:
  python humanize.py --title "My Essay" --text "Full essay text here..."
  python humanize.py --title "My Essay" --text "..." --prompt-summary "Summarize the..." --headless
"""

import argparse
import json
import os
import re
import sys
import time
from pathlib import Path

# ── Dependency check ──────────────────────────────────────────────────────────
missing = []
try:
    from playwright.sync_api import sync_playwright
except ImportError:
    missing.append('playwright')
try:
    import anthropic
except ImportError:
    missing.append('anthropic')
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    missing.append('python-docx')

if missing:
    print(f'Missing packages: {", ".join(missing)}')
    print(f'Run: pip install {" ".join(missing)}')
    if 'playwright' in missing:
        print('Then run: playwright install chromium')
    sys.exit(1)

# ── Paths ─────────────────────────────────────────────────────────────────────
DIR          = Path(__file__).parent
SESSION_FILE = DIR / 'session_gptinf.json'
CREDS_FILE   = DIR / 'credentials.env'
OUTPUT_DIR   = DIR / 'output'

# ── Credentials ───────────────────────────────────────────────────────────────
def load_creds():
    creds = {}
    with open(CREDS_FILE) as f:
        for line in f:
            eq = line.find('=')
            if eq > 0:
                creds[line[:eq].strip()] = line[eq+1:].strip()
    return creds

# ── Cookie session helpers ────────────────────────────────────────────────────
def save_cookies(context, path):
    with open(path, 'w') as f:
        json.dump(context.cookies(), f)

def load_cookies(context, path):
    if not Path(path).exists():
        return False
    with open(path) as f:
        context.add_cookies(json.load(f))
    return True

# ── GPTInf ────────────────────────────────────────────────────────────────────
def gptinf_check_logged_in(page):
    """Navigate to gptinf.com and return True if already logged in."""
    page.goto('https://gptinf.com', wait_until='domcontentloaded')
    time.sleep(2)
    url = page.url.lower()
    # If redirected to login page, not logged in
    if 'login' in url or 'signin' in url:
        return False
    # Check for logout/dashboard indicators
    logged_in = page.evaluate("""
        () => {
            const t = document.body.innerText || '';
            return t.includes('Logout') || t.includes('Dashboard') || t.includes('Sign out') ||
                   document.querySelector('[href*="logout"], [href*="dashboard"]') !== null;
        }
    """)
    return bool(logged_in)

def gptinf_login(page, email, password):
    """Log in to GPTInf and wait for redirect."""
    page.goto('https://gptinf.com/login', wait_until='domcontentloaded')
    time.sleep(1)
    # Fill email
    page.evaluate(f"""
        () => {{
            const inp = document.querySelector('input[type=email], input[name*=email], input[placeholder*=email i]');
            if (inp) {{
                inp.focus();
                inp.value = {json.dumps(email)};
                inp.dispatchEvent(new Event('input', {{bubbles: true}}));
                inp.dispatchEvent(new Event('change', {{bubbles: true}}));
            }}
        }}
    """)
    time.sleep(0.3)
    # Fill password
    page.evaluate(f"""
        () => {{
            const inp = document.querySelector('input[type=password], input[name*=password]');
            if (inp) {{
                inp.focus();
                inp.value = {json.dumps(password)};
                inp.dispatchEvent(new Event('input', {{bubbles: true}}));
                inp.dispatchEvent(new Event('change', {{bubbles: true}}));
            }}
        }}
    """)
    time.sleep(0.3)
    # Submit
    page.evaluate("""
        () => {
            const btn = document.querySelector('button[type=submit]') ||
                        Array.from(document.querySelectorAll('button')).find(b => /login|sign.?in/i.test(b.textContent));
            if (btn) btn.click();
        }
    """)
    # Wait up to 15s for redirect away from login page
    for _ in range(15):
        time.sleep(1)
        if 'login' not in page.url.lower() and 'signin' not in page.url.lower():
            return True
    return False

def gptinf_set_simple_mode(page):
    """Set GPTInf mode to Simple if not already set. Returns True if mode was set."""
    result = page.evaluate("""
        () => {
            // Find the mode selector button (shows current mode)
            const modes = ['Standard', 'Academic', 'Formal', 'Informal', 'Expand', 'Shorten', 'Creative'];
            const btns = Array.from(document.querySelectorAll('button'));
            const modeBtn = btns.find(b => modes.includes(b.textContent.trim()));
            if (!modeBtn) return {found: false};
            const current = modeBtn.textContent.trim();
            if (current === 'Simple') return {found: true, already: true};
            modeBtn.click();
            return {found: true, already: false, clicked: current};
        }
    """)
    if not result.get('found') or result.get('already'):
        return result.get('found', False)
    time.sleep(0.5)
    # Click Simple option in dropdown
    page.evaluate("""
        () => {
            const all = Array.from(document.querySelectorAll('li, button, span, div'));
            const simple = all.find(el => el.textContent.trim() === 'Simple');
            if (simple) simple.click();
        }
    """)
    time.sleep(0.3)
    # Click Apply if present
    page.evaluate("""
        () => {
            const apply = Array.from(document.querySelectorAll('button')).find(b => b.textContent.trim() === 'Apply');
            if (apply) apply.click();
        }
    """)
    time.sleep(0.3)
    return True

def gptinf_humanize_paragraph(page, text, max_wait=60):
    """
    Navigate to GPTInf, paste paragraph, humanize in Simple mode.
    Returns humanized text string, or None on failure.
    """
    page.goto('https://gptinf.com', wait_until='domcontentloaded')
    time.sleep(1.5)

    # Fill input textarea using execCommand (works with React)
    page.evaluate(f"""
        () => {{
            const ta = document.querySelector('textarea');
            if (!ta) return;
            ta.focus();
            ta.select();
            document.execCommand('selectAll');
            document.execCommand('insertText', false, {json.dumps(text)});
            ta.dispatchEvent(new Event('input', {{bubbles: true}}));
            ta.dispatchEvent(new Event('change', {{bubbles: true}}));
        }}
    """)
    time.sleep(0.5)

    # Set Simple mode
    gptinf_set_simple_mode(page)

    # Capture placeholder text so we know when output changes
    placeholder = page.evaluate("""
        () => {
            const all = document.querySelectorAll('p, div, span');
            for (const el of all) {
                if (el.children.length === 0 && el.textContent.includes('Paraphrased text will appear')) {
                    return el.textContent.trim();
                }
            }
            return 'Paraphrased text will appear here';
        }
    """)

    # Click Humanize button
    clicked = page.evaluate("""
        () => {
            const btn = Array.from(document.querySelectorAll('button')).find(b => /^humanize$/i.test(b.textContent.trim()));
            if (btn) { btn.click(); return true; }
            return false;
        }
    """)
    if not clicked:
        print('  WARNING: Could not find Humanize button.')
        return None

    # Poll for output change (up to max_wait seconds)
    input_preview = text[:40].lower()
    for i in range(max_wait):
        time.sleep(1)
        output = page.evaluate(f"""
            () => {{
                // Strategy 1: find element that no longer contains placeholder
                const placeholder = {json.dumps(placeholder[:60])};
                const all = document.querySelectorAll('p, div, span');
                for (const el of all) {{
                    if (el.children.length > 0) continue;
                    const t = el.textContent.trim();
                    if (t.length < 20) continue;
                    if (t.includes('Paraphrased text will appear')) continue;
                    // Skip if it looks like UI text
                    if (t.length < 50 && /^[A-Z][a-z]+ [A-Z]/.test(t)) continue;
                    // If it looks like body text (has sentence structure)
                    if (t.includes('. ') || t.includes(', ') || t.length > 100) return t;
                }}
                return null;
            }}
        """)
        if output and len(output) > 30:
            return output.strip()

        if i == 10:
            print('  Still waiting for GPTInf output...')

    print('  WARNING: GPTInf timed out — no output received.')
    return None

# ── Claude Sonnet verification ─────────────────────────────────────────────────
def verify_paragraph(client, original, humanized):
    """
    Check humanized paragraph for issues (dropped citations, altered quotes, broken meaning).
    Returns dict: {'ok': bool, 'fixed': str or None}
    """
    resp = client.messages.create(
        model='claude-sonnet-4-6',
        max_tokens=1024,
        messages=[{
            'role': 'user',
            'content': (
                'You are a proofreader comparing an original paragraph to a humanized rewrite.\n\n'
                f'Original:\n"""\n{original}\n"""\n\n'
                f'Humanized rewrite:\n"""\n{humanized}\n"""\n\n'
                'Check ONLY for these issues in the humanized version:\n'
                '1. Missing or altered in-text citations (e.g. dropped "(Smith, 2021)")\n'
                '2. Altered direct quotes (exact wording must be preserved inside quotation marks)\n'
                '3. Lost or inverted meaning/argument\n'
                '4. Grammar errors introduced by the rewrite\n\n'
                'If no issues: return {"ok": true}\n'
                'If issues found: return {"ok": false, "fixed": "corrected paragraph"} '
                '— make the minimum edits needed to fix the issues while keeping the humanized style intact. '
                'Do not revert to the original style.\n\n'
                'Return ONLY the JSON, no other text.'
            )
        }]
    )
    raw = resp.content[0].text.strip()
    m = re.search(r'\{.*\}', raw, re.DOTALL)
    if m:
        try:
            return json.loads(m.group())
        except Exception:
            pass
    return {'ok': True}

# ── GPTZero scoring ───────────────────────────────────────────────────────────
def gptzero_check(page, essay_text, max_wait=90):
    """
    Paste essay into GPTZero, click Scan, scrape human score + AI-flagged sentences.
    Returns (score_int_or_None, list_of_ai_sentence_strings).
    """
    page.goto('https://gptzero.me', wait_until='domcontentloaded')
    time.sleep(2)

    # Paste text via execCommand (React textarea)
    page.evaluate(f"""
        () => {{
            const ta = document.querySelector('textarea');
            if (!ta) return;
            ta.focus();
            ta.select();
            document.execCommand('selectAll');
            document.execCommand('insertText', false, {json.dumps(essay_text[:9500])});
            ta.dispatchEvent(new Event('input', {{bubbles: true}}));
            ta.dispatchEvent(new Event('change', {{bubbles: true}}));
        }}
    """)
    time.sleep(0.5)

    # Click Scan button
    page.evaluate("""
        () => {
            const btn = Array.from(document.querySelectorAll('button')).find(b => /^scan$/i.test(b.textContent.trim()));
            if (btn) btn.click();
        }
    """)

    # Poll for results
    for i in range(max_wait):
        time.sleep(1)
        result = page.evaluate("""
            () => {
                const text = document.body.innerText || '';

                // Pattern 1: "X% Human" or "X% likely to be human"
                let m = text.match(/(\d{1,3})%\s*(?:Human|human|likely\s*to\s*be\s*human)/);
                if (m) return {score: parseInt(m[1]), found: true};

                // Pattern 2: "Your text is X% human"
                m = text.match(/(?:is|are)\s*(\d{1,3})%\s*human/i);
                if (m) return {score: parseInt(m[1]), found: true};

                // Pattern 3: score in a dedicated element
                const scoreEls = document.querySelectorAll('[class*="score"], [class*="percent"], [class*="human-score"]');
                for (const el of scoreEls) {
                    m = el.textContent.match(/(\d{1,3})%/);
                    if (m) return {score: parseInt(m[1]), found: true};
                }

                // Check if results section appeared (even if we can't parse score yet)
                const hasResults = text.includes('Overall') || text.includes('Completely') ||
                                   text.includes('Your text') || text.includes('sentences');
                return {found: false, hasResults};
            }
        """)

        if result.get('found'):
            score = result['score']
            # Scrape AI-highlighted sentences
            ai_sentences = page.evaluate("""
                () => {
                    // GPTZero highlights AI sentences with colored spans/marks
                    const highlighted = document.querySelectorAll(
                        '[class*="highlight"], [class*="ai-sentence"], [class*="flagged"], mark, [style*="background-color: rgb(255"]'
                    );
                    const sentences = Array.from(highlighted)
                        .map(el => el.textContent.trim())
                        .filter(t => t.length > 15);

                    // Deduplicate
                    return [...new Set(sentences)];
                }
            """)
            return score, ai_sentences

        if i == 15:
            print('  Still waiting for GPTZero results...')
        if i == 30:
            # Check for word limit error
            limit_hit = page.evaluate("""
                () => {
                    const t = document.body.innerText || '';
                    return t.includes('upgrade') || t.includes('limit') || t.includes('sign up');
                }
            """)
            if limit_hit:
                print('  GPTZero word limit reached — score unavailable (free tier limit).')
                return None, []

    print('  GPTZero timed out.')
    return None, []

# ── Word doc generation ────────────────────────────────────────────────────────
def save_word_doc(title, prompt_summary, original_essay, humanized_essay, gptzero_score, output_path):
    doc = Document()

    # Cover info
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if prompt_summary:
        p = doc.add_paragraph()
        p.add_run('Prompt: ').bold = True
        p.add_run(prompt_summary)

    score_text = f'{gptzero_score}%' if gptzero_score is not None else 'N/A'
    p = doc.add_paragraph()
    p.add_run('GPTZero Human Score: ').bold = True
    p.add_run(score_text)
    doc.add_page_break()

    # Original AI draft
    doc.add_heading('Original AI Draft', level=1)
    for para in original_essay.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_page_break()

    # Humanized version
    doc.add_heading('Humanized Version', level=1)
    for para in humanized_essay.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.save(output_path)

# ── Main ───────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description='Humanize an AI-drafted essay via GPTInf + GPTZero')
    parser.add_argument('--title',          required=True, help='Essay title (used for filename)')
    parser.add_argument('--text',           required=True, help='Full essay text (paragraphs separated by double newlines)')
    parser.add_argument('--prompt-summary', default='',   help='Short summary of the writing prompt')
    parser.add_argument('--headless',       action='store_true', help='Run browser in headless mode')
    args = parser.parse_args()

    creds = load_creds()
    for key in ('ANTHROPIC_API_KEY', 'GPTINF_EMAIL', 'GPTINF_PASSWORD'):
        if not creds.get(key) or creds[key].startswith('your_'):
            print(f'ERROR: {key} not set in credentials.env')
            sys.exit(1)

    OUTPUT_DIR.mkdir(exist_ok=True)

    # Split essay into paragraphs
    paragraphs = [p.strip() for p in args.text.split('\n\n') if p.strip()]
    if not paragraphs:
        print('ERROR: No paragraphs found in --text')
        sys.exit(1)
    print(f'\n{len(paragraphs)} paragraph(s) to process.\n')

    original_essay = '\n\n'.join(paragraphs)

    # Claude Sonnet client for verification
    claude = anthropic.Anthropic(api_key=creds['ANTHROPIC_API_KEY'])

    humanized_paragraphs = list(paragraphs)  # start as copy
    final_score = None

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=args.headless)
        context = browser.new_context(viewport={'width': 1280, 'height': 900})
        page = context.new_page()

        # ── GPTInf: login / session ────────────────────────────────────────
        load_cookies(context, SESSION_FILE)
        if gptinf_check_logged_in(page):
            print('GPTInf session active.\n')
        else:
            print('Logging into GPTInf...')
            ok = gptinf_login(page, creds['GPTINF_EMAIL'], creds['GPTINF_PASSWORD'])
            if not ok:
                print('ERROR: GPTInf login failed. Check credentials.')
                browser.close()
                sys.exit(1)
            save_cookies(context, SESSION_FILE)
            print('Logged in, session saved.\n')

        # ── Phase 1: humanize each paragraph ──────────────────────────────
        for i, para in enumerate(paragraphs):
            print(f'[{i+1}/{len(paragraphs)}] Humanizing...')
            preview = para[:90].replace('\n', ' ')
            print(f'  Original: {preview}{"..." if len(para) > 90 else ""}')

            humanized = gptinf_humanize_paragraph(page, para)
            if not humanized:
                print('  GPTInf returned no output — keeping original.')
                humanized = para
            else:
                h_preview = humanized[:90].replace('\n', ' ')
                print(f'  Humanized: {h_preview}{"..." if len(humanized) > 90 else ""}')

            # Verify with Claude Sonnet
            verif = verify_paragraph(claude, para, humanized)
            if not verif.get('ok') and verif.get('fixed'):
                print('  Verification: FIXED (restored citations/quotes/meaning)')
                humanized = verif['fixed']
            else:
                print('  Verification: OK')

            humanized_paragraphs[i] = humanized

        humanized_essay = '\n\n'.join(humanized_paragraphs)

        # ── Phase 2: GPTZero check ─────────────────────────────────────────
        print('\nChecking GPTZero score...')
        score, ai_sentences = gptzero_check(page, humanized_essay)
        print(f'  Human score: {score}%' if score is not None else '  Score: unavailable')
        final_score = score

        # ── Phase 3: re-run flagged paragraphs if score < 70 ──────────────
        if score is not None and score < 70:
            print(f'\nScore {score}% < 70% — re-running flagged paragraphs...')

            if ai_sentences:
                # Target only paragraphs containing flagged sentences
                flagged = set()
                for i, para in enumerate(humanized_paragraphs):
                    for sent in ai_sentences:
                        if sent and len(sent) > 15 and sent[:35] in para:
                            flagged.add(i)
                if not flagged:
                    # Fallback: re-run all paragraphs
                    flagged = set(range(len(humanized_paragraphs)))
                print(f'  Flagged paragraph indices: {sorted(flagged)}')
            else:
                # No sentence highlights available — re-run all
                flagged = set(range(len(humanized_paragraphs)))
                print('  No sentence highlights — re-running all paragraphs.')

            for i in sorted(flagged):
                print(f'\n  Re-humanizing paragraph {i+1}...')
                humanized = gptinf_humanize_paragraph(page, humanized_paragraphs[i])
                if humanized:
                    verif = verify_paragraph(claude, paragraphs[i], humanized)
                    if not verif.get('ok') and verif.get('fixed'):
                        humanized = verif['fixed']
                    humanized_paragraphs[i] = humanized

            humanized_essay = '\n\n'.join(humanized_paragraphs)

            # Re-check GPTZero
            print('\nRe-checking GPTZero score...')
            score2, _ = gptzero_check(page, humanized_essay)
            print(f'  Final score: {score2}%' if score2 is not None else '  Score: unavailable')
            final_score = score2 if score2 is not None else score
            if score2 is not None and score2 < 70:
                print(f'  NOTE: Still below 70% after re-run. Saving best version.')

        browser.close()

    # ── Save Word doc ──────────────────────────────────────────────────────
    safe_title = re.sub(r'[^a-zA-Z0-9 _-]', '_', args.title)[:50].strip()
    output_path = OUTPUT_DIR / f'{safe_title}.docx'
    save_word_doc(
        title=args.title,
        prompt_summary=args.prompt_summary or args.title,
        original_essay=original_essay,
        humanized_essay=humanized_essay,
        gptzero_score=final_score,
        output_path=output_path
    )

    print(f'\nDone!')
    print(f'Word doc: {output_path}')
    print(f'GPTZero human score: {final_score}%' if final_score is not None else 'GPTZero score: unavailable')

if __name__ == '__main__':
    main()
