"""
Word document writer using python-docx.
Saves answer keys and study guides as .docx files locally.
"""

from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def write_answer_key(
    output_dir: str | Path,
    course_name: str,
    assignment_name: str,
    questions_answers: list[dict],
    study_guide: str,
) -> Path:
    """
    Creates (or overwrites) a .docx file with the answer key and study guide.
    Returns the path to the saved file.
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Sanitize filename
    safe_name = "".join(c if c.isalnum() or c in " -_" else "_" for c in assignment_name)
    filename = f"{course_name} - {safe_name}.docx"
    filepath = output_dir / filename

    doc = Document()

    # --- Title ---
    title = doc.add_heading(f"{course_name}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_heading(assignment_name, level=2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    now = datetime.now().strftime("%B %d, %Y  %I:%M %p")
    ts = doc.add_paragraph(f"Generated: {now}")
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # --- Answer Key section ---
    doc.add_heading("Answer Key", level=2)
    doc.add_paragraph()

    for i, qa in enumerate(questions_answers, 1):
        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{i}. ")
        q_run.bold = True
        q_para.add_run(qa["question"])

        # Answer
        a_para = doc.add_paragraph()
        a_para.paragraph_format.left_indent = Pt(18)
        a_run = a_para.add_run("Answer: ")
        a_run.bold = True
        a_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        a_para.add_run(qa["correct_answer"])

        # Explanation
        if qa.get("explanation"):
            e_para = doc.add_paragraph()
            e_para.paragraph_format.left_indent = Pt(18)
            e_run = e_para.add_run("Why: ")
            e_run.italic = True
            e_para.add_run(qa["explanation"])

        doc.add_paragraph()  # spacer between questions

    # --- Study Guide section ---
    doc.add_page_break()
    doc.add_heading("Study Guide", level=2)
    doc.add_paragraph()

    # Write study guide line by line, preserving headers (lines starting with #)
    for line in study_guide.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=3)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    doc.save(filepath)
    return filepath
